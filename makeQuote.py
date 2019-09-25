#! Python 3
# This code generates a polished JCI quote from simple prompts... speeding up quote generation time.
# This code makes use of Python-docx v0.8.7.

def makeQuote(excelQuote):
    #import libraries
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
    from docx.enum.table import WD_ALIGN_VERTICAL
    import datetime, os, openpyxl
    from quoteFunctions import set_col_widths, getScopeList, getClarifyList, updateTracker
    from docx.shared import Inches
    from TandCs import TandCs

# Open a new workbook
    wb = openpyxl.load_workbook(excelQuote)
    sheet = wb["Sheet1"]

    # Get all needed information so we can write it into a word document:
    #quoteID = str(input("NxGen Quote ID (1-XXXXXXX): "))
    quoteID = sheet["B1"].value
    #customer = str(input("Customer: "))
    customer = sheet["B2"].value
    #contact = str(input("Quote POC: "))
    contact = sheet["B3"].value
    #projectName = str(input("Name of the project? "))
    projectName = sheet["B4"].value
    #price = str(input("Project price? "))
    price = sheet["B5"].value
    time_stamp = str(datetime.date.today().strftime("%Y%m%d"))
    jciAddress = "Washington DC Service Branch\n1101 Hampton Park Blvd. Bldg. C Ste 100\n Capitol Heights, MD 20743\n301-324-4500\n"
    scopeIntro = "Johnson Controls, Inc. (JCI) is pleased to provide the following scope of work and pricing for the project described below."
    outputFile = (f'{customer} - {projectName} - {time_stamp}.docx')
    xlNumber = (f'{customer} - {projectName} - {time_stamp}.xlsx')

    # Open Word Document Template,
    doc = Document('Quote.docx')

    #Initialize Header With Quote ID & Branch Address
    doc.add_heading('Quote: ' + quoteID, 1)
    address = doc.add_paragraph(jciAddress)
    address.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    #Insert Header Table
    table = doc.add_table(2,4,None)
    header = ["To:", customer, "Date:", datetime.date.today().strftime("%A, %B %d, %Y"),"Attn:", contact, "Project:", projectName]
    i = 0
    for row in table.rows:
        for cell in row.cells:
            cell.text = header[i]
            if i % 2 == 0:
                for paragraph in cell.paragraphs:
                    paragraph.style = "Table Header"
            i += 1
    set_col_widths(table, 0.1, 0.61, 1.6, .71, 3.34) 

    #Populate the Scope of Work Section of the Quote
    doc.add_paragraph(f'\n{scopeIntro}\n')
    doc.add_paragraph("Scope of Work:", "Quote Section")

    scopeList = getScopeList(excelFile)
    cellStart = scopeList.pop()

    for item in scopeList:
        doc.add_paragraph(item, "Numbered List")   

    #Project pricing
    doc.add_paragraph("\nProject Pricing:", "Quote Section")
    p = doc.add_paragraph(f"Pricing for this project is: ", "Bullet List")
    p.add_run(f'{price}').bold = True
    p.add_run('.').bold = False

    #Clarifications and Exclusions
    doc.add_paragraph("\nClarifications and Exclusions", "Quote Section")
    p = doc.add_paragraph("All work will be accomplished during normal business hours", "Bullet List")
    p = doc.add_paragraph("Asbestos and Lead remediation is excluded", "Bullet List")
    p = doc.add_paragraph("Any work not explicitly outlined in the 'Scope of Work' section above is excluded", "Bullet List")
    p = doc.add_paragraph("Drywall, painting, and insulation work is excluded unless specified in the scope of work", "Bullet List")
    p = doc.add_paragraph("Johnson Controls reserves the right to progress bill for this work", "Bullet List")

    clarifyList = getClarifyList(cellStart, excelFile)

    for item in clarifyList:
        doc.add_paragraph(item, "Bullet List")

    #add a signature block:
    p = doc.add_paragraph("\nPlease don't hesitate to call me at 202.424.9705 if you have any questions.  Thank you for working with Johnson Controls, Inc.!")
    p.paragraph_format.keep_with_next = True
    p = doc.add_paragraph("~", "Jim Signature")
    p.paragraph_format.left_indent = Inches(3.1)
    p.paragraph_format.keep_with_next = True
    signatureBlock = ["Jim Bergren", "Account Representative, Washington, DC", "Johnson Controls, Inc."]

    for item in signatureBlock:
        p = doc.add_paragraph(item)
        p.paragraph_format.left_indent = Inches(3.25)
        p.paragraph_format.keep_with_next = True

    p = doc.add_paragraph()
    run = p.add_run()
    run.add_break(WD_BREAK.PAGE)

    expir = ["This proposal and alternates listed below are hereby accepted and Johnson Controls is authorized to proceed with work; subject, however to credit approval by Johnson Controls, Inc., Milwaukee, Wisconsin.","This proposal is valid for 30 days"]
    table = doc.add_table(1,2, None)
    i = 0
    for row in table.rows:
        for cell in row.cells:
            cell.text = expir[i]
            for paragraph in cell.paragraphs:
                paragraph.style = "Small Print"
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            i += 1
    doc.add_paragraph()

    approvalBlock = [
        "", customer, "", "Johnson Controls, Inc.", 
        "Name:", "______________________________________", "Name:", "______________________________________", 
        "Title:", "______________________________________", "Title:", "______________________________________", 
        "Date:", "______________________________________","Date:", "______________________________________",
        "PO #:","______________________________________","",""
        ]

    table = doc.add_table(5,4, None)
    i = 0
    for row in table.rows:
        for cell in row.cells:
            cell.text = approvalBlock[i]
            if i < 4:
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            else:
                cell.vertical_alignment = WD_ALIGN_VERTICAL.BOTTOM

            for paragraph in cell.paragraphs:
                if i <= 3:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                else:
                    paragraph.style = "Table Header"
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            i += 1
    set_col_widths(table, 0.4, 0.67, 2.83, 0.67, 2.83)

    doc.add_paragraph('\n')



    #Terms and Conditions
    pageBreak = doc.add_paragraph()
    run = pageBreak.add_run()
    run.add_break(WD_BREAK.PAGE)

    doc.add_paragraph("Terms and Conditions", "Quote Section")

    for each in TandCs:
        p = doc.add_paragraph(each, "Terms and Conditions")

    #Save the Files
    outputPath = os.path.join("C:\\users\\jbergrj\\OneDrive - Johnson Controls\\Quotes")
    os.chdir(outputPath)
    doc.save(outputFile)

    xlOutputPath = os.path.join("C:\\users\\jbergrj\\OneDrive - Johnson Controls\\Quotes\\Quote Numbers")
    os.chdir(xlOutputPath)
    wb.save(xlNumber)

    #populate L&M Tracker
    updateTracker(quoteID, projectName, customer, contact, price)

    print ("Quote Generated")
excelFile = str(input("Enter an excel filename to generate a quote: "))
makeQuote(excelFile)
