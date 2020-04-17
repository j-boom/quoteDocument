#! Python 3

import datetime, os

from docx import Document
from docx.shared import Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_ALIGN_VERTICAL
import openpyxl

class Quote:
    def __init__(self):
        self.quote_number = 'sheet["B1"].value'
        self.point_of_contact = 'sheet["B2"].value'
        self.company_name = 'sheet["B3"].value'
        self.company_address_one = 'sheet["B4"].value'
        self.company_address_two = 'sheet["B5"].value'
        self.project_site = 'sheet["B6"].value'
        self.project_name = 'sheet["B7"].value'
        self.price = 'sheet["B8"].value'

        self.pretty_date = str(datetime.date.today().strftime("%A, %B %d, %Y"))
        self.time_stamp = str(datetime.date.today().strftime("%Y%m%d"))
        self.expiration_date = (datetime.date.today() + datetime.timedelta(days=90)).strftime("%B %d, %Y")


def make_quote(excel_input, scope_list, exclusions_list, clarifications_list, rep_name, rep_email, t_cs, quote):   
    doc = Document('C:\\Users\\jbergren\\OneDrive - ENGINEERED SERVICES INC\\Quotes\\quote_template.docx')
    print('...Inserting Project Information')
    doc.add_paragraph('{}'.format(quote.pretty_date), style="ESI Date")
    doc.add_paragraph('{}'.format(quote.quote_number), style="ESI Date")
    doc.add_heading('PROPOSAL', 1)
    doc.add_heading('Attn: {}'.format(quote.point_of_contact), 3)
    doc.add_heading('{}'.format(quote.company_name), 3)
    doc.add_paragraph('{}'.format(quote.company_address_one))
    doc.add_paragraph('{}'.format(quote.company_address_two))
    doc.add_paragraph('')
    doc.add_heading('{} - {}'.format(quote.project_site, quote.project_name), 2)
    doc.add_paragraph('')

    print('...Inserting Scope of Work')
    doc.add_heading("Scope of Work:", 3)

    for item in scope_list:
        doc.add_paragraph(item, "ScopeOfWork")   

    print('...Getting Exclusions')
    doc.add_heading("Exclusions:", 4)

    if len(exclusions_list) > 0:
        for item in exclusions_list:
            doc.add_paragraph(item, "ClarificationsExclusionsBold")
    doc.add_paragraph("Any overtime or off-hours work.", "ClarificationsExclusionsNormal")
    doc.add_paragraph("3rd party or additional commissioning", "ClarificationsExclusionsNormal")
    doc.add_paragraph("Engineering drawings for construction or as-builts", "ClarificationsExclusionsNormal")
    doc.add_paragraph("Drywall repair or patching, painting, and insulation work unless specified in the scope of work", "ClarificationsExclusionsNormal")
    doc.add_paragraph("Any work not explicitly outlined in the 'Scope of Work' section above", "ClarificationsExclusionsNormal")
    doc.add_paragraph("")

    print('...Clarifying Things')
    doc.add_heading("Clarifications:", 4)

    if len(clarifications_list) > 0:
        for item in clarifications_list:
            doc.add_paragraph(item, "ClarificationsExclusionsBold")
    doc.add_paragraph("Standard one year parts and labor warrantly on all new equipment is included.", "ClarificationsExclusionsNormal")
    doc.add_paragraph("")

    print('...Applying Pricing')
    p = doc.add_paragraph("Our price for the above work: ", "Heading 3")
    p.add_run('{}'.format(quote.price)).underline = True
    p.add_run('.').underline = False
    doc.add_paragraph("")

    print('...Signing for you')
    signatureBlock = ["{}, Account Manager".format(rep_name), "Engineered Services, Inc.", "{}".format(rep_email)]

    doc.add_paragraph("~", "Rep_Signature")
    
    for item in signatureBlock:
        doc.add_paragraph(item, "Signature Block")

    doc.add_paragraph("")
    doc.add_paragraph("The price quoted above is guaranteed until {}.  After this date, we may require re-pricing and/or re-scheduling the work.\n".format(quote.expiration_date))
        

    approval_block = [
        "{}\nAccepetance".format(quote.company_name), "This proposal and alternates listed below are hereby accepted\n and ESI is authorized to proceed with work.\n Subject, however to credit approval by ESI",
        "Signature:", "_________________________________________________________",
        "Name:", "_________________________________________________________", 
        "Title:", "_________________________________________________________", 
        "Date:", "_________________________________________________________",
        "PO #:","_________________________________________________________"
        ]

    table = doc.add_table(6,2)
    table.style = 'Approval'
    i = 0
    for row in table.rows:
        for cell in row.cells:
            cell.text = approval_block[i]
        
            for paragraph in cell.paragraphs:
                if i == 0:
                    paragraph.style = "Heading 3"
                elif i == 1:
                    paragraph.style = "FinePrintItal"
                elif i > 1:  
                    cell.vertical_alignment = WD_ALIGN_VERTICAL.BOTTOM  
                    if i % 2 == 0 and i > 1:
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    else: 
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            i += 1

    widths = [Inches(.71), Inches(3.34)]
    for row in table.rows:
        row.height = Inches(0.4)
    for i, width in enumerate(widths):
        row.cells[i].width = width

    doc.add_paragraph('\n')
 
    print('...Adding the fine print')
    pageBreak = doc.add_paragraph()
    run = pageBreak.add_run()
    run.add_break(WD_BREAK.PAGE)

    doc.add_heading("STANDARD CONDITIONS", 2)

    for each in t_cs:
        doc.add_paragraph(each, "FinePrint")

    output_file = '{} - {} - {} - {}.docx'.format(quote.company_name, quote.project_site, quote.project_name, quote.time_stamp)
    print('...Saving {}'.format(output_file))
    proposal_output_path = os.path.join("C:\\Users\\jbergren\\OneDrive - ENGINEERED SERVICES INC\\Quotes")
    os.chdir(proposal_output_path)
    doc.save(output_file)